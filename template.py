from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches
import datetime
import json
from docx2pdf import convert


def date(dt):
    dtl = dt.split("/")
    d = int(dtl[1])
    m = int(dtl[0])
    y = int(dtl[2])+2000
    x = datetime.datetime(y, m, d)
    return x


with open("data.json", "r") as d:
    data = json.load(d)

for i in data:
    numd = list(i.keys())[0]
    info = i[numd]

    date_E = date(info["date_E"]).strftime("%d,%b,%Y")
    try:
        date_S = date(info["date_S"]).strftime("%d,%b,%Y")
    except:
        date_S = info["date_S"]
    document = Document()

    document.add_heading('Fiche Technique Pancreatite aigue lithiasique', 0)

    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Num Dossier : ").bold = True
    p.add_run(numd)

    p = document.add_paragraph()
    p.add_run("Nom : ").bold = True
    p.add_run(info["nom"])

    p = document.add_paragraph()
    p.add_run("Prénom : ").bold = True
    p.add_run(info["prenom"])

    p = document.add_paragraph()
    p.add_run("Date d'entré : ").bold = True
    p.add_run(str(date_E))

    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Date du sortie : ").bold = True
    p.add_run(str(date_S))

    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Durée du sejour hospitalier : ").bold = True
    if info["sejour"] == "Inconnu":
        p.add_run(str(info["sejour"]))
    else:
        p.add_run(str(info["sejour"]) + " Jours")

    p = document.add_paragraph()
    p.add_run("Age: ").bold = True
    p.add_run(str(info["age"]) + " ans")

    p = document.add_paragraph()
    p.add_run("Sex: ").bold = True
    if info["sex"]:
        p.add_run("Male")
    else:
        p.add_run("Female")

    p = document.add_paragraph()
    p.add_run("IMC: ").bold = True
    if not info["imc"]:
        p.add_run("Pas calculé")
    else:
        p.add_run(str(info["imc"]))

    document.add_heading('Motifs de consultation', level=1)
    for i in info["motifs"]:
        document.add_paragraph(i, style='List Bullet')

    document.add_heading('Tableau clinique', level=1)
    # ☑
    document.add_heading('Signes Généraux', level=2)
    p = document.add_paragraph()
    if info["fievre"]:
        p.add_run("   ☑ Fievre     ").bold = True
    else:
        p.add_run("   ☐ Fievre     ").bold = True
    if info["asthenie"]:
        p.add_run("   ☑ Asthénie     ").bold = True
    else:
        p.add_run("   ☐ Asthénie     ").bold = True
    if info["aeg"]:
        p.add_run("   ☑ AEG  ").bold = True
    else:
        p.add_run("   ☐ AEG  ").bold = True
    document.add_heading('Signes fonctionneles', level=2)
    p = document.add_paragraph()
    p.add_run("Douleur Typique: ").bold = True
    if info["dlr"]:
        p.add_run(" ☑ OUI         ☐ NON")
    else:
        p.add_run(" ☐ OUI         ☑ NON")
    p = document.add_paragraph()
    p.add_run("Vomissement: ").bold = True
    if info["vom"]:
        p.add_run(" ☑ OUI         ☐ NON")
    else:
        p.add_run(" ☐ OUI         ☑ NON")
    document.add_heading('Signes physiques : Normal?', level=2)
    p = document.add_paragraph()

    if info["sps"] == []:
        p.add_run(" ☑ OUI         ☐ NON").bold = True
    else:
        p.add_run(" ☐ OUI         ☑ NON").bold = True
        for i in info["sps"]:
            document.add_paragraph(i, style='List Bullet')
    document.add_heading('SIRS', level=2)
    p = document.add_paragraph()
    if info["sirs"]:
        p.add_run(" ☑ POSITIF         ☐ NEGATIF")
    else:
        p.add_run(" ☐ POSITIF         ☑ NEGATIF")
    document.add_heading('Examens biologique', level=1)
    document.add_heading('FNS', level=2)
    p = document.add_paragraph()
    p.add_run("        Gb : "+str(info["gb"]) + " x10³/mm³").bold = True
    p.add_run("        Hb : "+str(info["hb"]) + " g/dl").bold = True
    p.add_run("        Plq : "+str(info["plq"]) + " x10³/mm³").bold = True
    document.add_heading(
        'Lipasémie : '+str(info["lip"]) + "x  normal", level=2)
    document.add_heading('CRP', level=2)
    p = document.add_paragraph()
    if info["crp"]:
        p.add_run(" ☑ POSITIF         ☐ NEGATIF")
    else:
        p.add_run(" ☐ POSITIF         ☑ NEGATIF")
    if not info["bio_autres"] == {}:
        document.add_heading('Autres: ', level=2)
        for i in info["bio_autres"]:
            document.add_paragraph(i, style='List Bullet')

    document.add_heading('Examens radiologiques', level=1)
    document.add_heading('Echographie abdomino-pelvienne', level=2)
    p = document.add_paragraph()
    p.add_run(info["echo"])
    document.add_heading(
        'TDM avec ou sans injection : Score Baltazard [ ' + info["balt"] + " ]", level=2)
    document.add_heading('Conduite thérapeutique ', level=1)
    document.add_heading('Mésures de réanimation', level=2)
    p = document.add_paragraph()
    if info["rea"]:

        p.add_run(" ☑ OUI         ☐ NON")
    else:
        p.add_run(" ☐ OUI         ☑ NON")
    document.add_heading('Antalgique', level=2)
    p = document.add_paragraph()
    p.add_run(" Palier : " + str(info["anta"])).bold = True
    if not info["atbq"] == []:
        document.add_heading('Antibiothérapie', level=2)
        p = document.add_paragraph()
        p.add_run(" Fievre").bold = True
        if info["atbq"][0]:
            p.add_run(" ☑ OUI         ☐ NON")
        else:
            p.add_run(" ☐ OUI         ☑ NON")
        p = document.add_paragraph()
        p.add_run(" Elevation du Gb").bold = True
        if info["atbq"][1]:
            p.add_run(" ☑ OUI         ☐ NON")
        else:
            p.add_run(" ☐ OUI         ☑ NON")
    document.add_heading('Evolution', level=2)
    p = document.add_paragraph()

    if info["evod"] == []:
        p.add_run(" ☑ Favorable         ☐ Défavorable").bold = True
    else:
        p.add_run(" ☐ Favorable         ☑ Défavorable").bold = True
        for i in info["evod"]:
            document.add_paragraph(i, style='List Bullet')
    document.add_heading('Treatement', level=2)
    p = document.add_paragraph()
    p.add_run(" Médical   ").bold = True
    if info["trtmed"]:

        p.add_run(" ☑ OUI         ☐ NON")
    else:
        p.add_run(" ☐ OUI         ☑ NON")
    p = document.add_paragraph()
    p.add_run(" Chirurgicale   ").bold = True
    if info["trtchir"]:

        p.add_run(" ☑ OUI         ☐ NON")
    else:
        p.add_run(" ☐ OUI         ☑ NON")
    if not info["moment"] == "none":
        p = document.add_paragraph()
        p.add_run(" Moment d'intervention ").bold = True
        if info["moment"]:
            p = document.add_paragraph()
            p.add_run(" ☑ Au cours du meme sejour ")
            p.add_run(" ☐ 6 ou 8eme Semaine apres ")
        else:
            p = document.add_paragraph()
            p.add_run(" ☐ Au cours du meme sejour ")
            p.add_run(" ☑ 6 ou 8eme Semaine apres ")
    if not info["pourquoi"] == "":
        p = document.add_paragraph()
        for i in info["pourquoi"]:
            document.add_paragraph(i, style='List Bullet')

    document.add_heading('Suites post hospitaliers', level=1)
    p = document.add_paragraph()
    if info["suites"]:

        p.add_run(" ☑ Simples         ☐ Compliquees")
    else:
        p.add_run(" ☐ Simples         ☑ Compliquees")

    docname = info["nom"] + " " + info["prenom"]
    document.save("docx/"+docname+".docx")
    #convert("docx/"+docname+".docx", "docx/"+docname+".pdf")
